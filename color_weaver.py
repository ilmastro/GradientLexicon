from docx import Document
from docx.shared import Pt, RGBColor

def hex_to_rgb(hex_color):
    # Convert a hex color to an RGB tuple
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def read_text_from_file(file_path):
    # Read text from a given file with UTF-8 encoding
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()


def apply_colors_to_lines(text, color_palette, output_file, font_name, font_size):
    # Create a new Word document
    doc = Document()
    
    # Split the text into lines
    lines = text.split('\n')

    # Apply colors to each line
    color_index = 0
    for line in lines:
        # Create a new paragraph for each line
        paragraph = doc.add_paragraph()
        
        # Split line into words for coloring
        words = line.split()
        for word in words:
            color = hex_to_rgb(color_palette[color_index % len(color_palette)])
            run = paragraph.add_run(word + ' ')
            run.font.color.rgb = RGBColor(color[0], color[1], color[2])
            run.font.name = font_name
            run.font.size = Pt(font_size)
        
        color_index += 1

    # Save the document
    doc.save(output_file)

def split_poetry_lines(file_path):
    with open(file_path, 'r') as file:
        contents = file.read()
    lines = contents.split('\n')
    return [line for line in lines if line.strip() != '']

# Define palette of colors
blue_palette = ["4da4b9","4aa6bb","4ca7bc","4da8bd","4ea8bd","4ea8bd","4fa9be","50a9be","51aabe","52aabe"]

# Path to text file
text_file_path = "GradientLexicon/wit.txt"

# Read the text from the file
text = read_text_from_file(text_file_path)

# Apply colors to each line of the poetry
apply_colors_to_lines(text, blue_palette, "GradientLexicon/output.docx", "Source Sans 3", 16)
